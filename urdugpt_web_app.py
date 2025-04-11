import streamlit as st
import torch
import torch.nn.functional as F
from tokenizers import Tokenizer
from urdugpt_step8_transformer import Transformer
from urdugpt_utils import load_config
from urdugpt_step2_dataloader import causal_mask
import os
import pandas as pd
from io import BytesIO
from docx import Document
from fpdf import FPDF

# Load config
config = load_config()
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")

max_seq_len = int(open(config['data']['max_seq_len_path']).read().strip())
tokenizer_en = Tokenizer.from_file(config['data']['tokenizer_en_path'])
tokenizer_ur = Tokenizer.from_file(config['data']['tokenizer_ur_path'])

# Load model
def load_model():
    model = Transformer(
        tokenizer_en.get_vocab_size(),
        tokenizer_ur.get_vocab_size(),
        max_seq_len, max_seq_len,
        d_model=config['model']['d_model'],
        num_layers=config['model']['num_layers'],
        num_heads=config['model']['num_heads'],
        d_ff=config['model']['d_ff'],
        dropout_rate=config['model']['dropout']
    ).to(device)
    ckpt_path = os.path.join(config['training']['checkpoint_dir'], config['training']['checkpoint_name'])
    model.load_state_dict(torch.load(ckpt_path, map_location=device))
    model.eval()
    return model

@st.cache_data
def convert_df_to_csv(df):
    return df.to_csv(index=False).encode("utf-8")

def export_to_pdf(data):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for eng, urdu in data:
        pdf.multi_cell(0, 10, f"English: {eng}\nUrduGPT: {urdu}\n", align="L")
    return pdf.output(dest='S').encode('latin1')

def export_to_docx(data):
    doc = Document()
    doc.add_heading("UrduGPT Translations", 0)
    for eng, urdu in data:
        doc.add_paragraph(f"English: {eng}")
        doc.add_paragraph(f"UrduGPT: {urdu}", style="Intense Quote")
    output = BytesIO()
    doc.save(output)
    return output.getvalue()

# Translation with token-level confidence scores
def decode_with_confidence(model, text):
    tokens = tokenizer_en.encode(text).ids
    cls = tokenizer_en.token_to_id("[CLS]")
    sep = tokenizer_en.token_to_id("[SEP]")
    pad = tokenizer_en.token_to_id("[PAD]")
    src = torch.tensor([cls] + tokens + [sep] + [pad]*(max_seq_len - len(tokens) - 2), dtype=torch.int64).unsqueeze(0).to(device)
    src_mask = (src != pad).unsqueeze(0).unsqueeze(0).int()

    output_ids = []
    output_tokens = []
    output_conf = []
    decoder_input = torch.tensor([[tokenizer_ur.token_to_id("[CLS]")]], device=device)
    encoder_output = model.encode(src, src_mask)

    for _ in range(max_seq_len):
        dec_mask = causal_mask(decoder_input.size(1)).type_as(src_mask).to(device)
        out = model.decode(decoder_input, encoder_output, src_mask, dec_mask)
        logits = model.projection(out[:, -1])
        probs = F.softmax(logits, dim=-1)
        conf, next_token = torch.max(probs, dim=-1)

        token_id = next_token.item()
        confidence = conf.item()
        output_ids.append(token_id)
        output_tokens.append(tokenizer_ur.id_to_token(token_id))
        output_conf.append(confidence)

        decoder_input = torch.cat([decoder_input, next_token.unsqueeze(0)], dim=1)
        if token_id == tokenizer_ur.token_to_id("[SEP]"):
            break

    return output_ids, output_tokens, output_conf

# UI START
st.set_page_config(
    page_title="UrduGPT Translator",
    page_icon="./favicon.ico",
    layout="centered"
)

# Custom CSS styling
st.markdown("""
<style>
body {
    background-color: #f8f9fa;
    font-family: 'Segoe UI', sans-serif;
}
.css-18e3th9 {
    background-color: #ffffff;
    padding: 2rem;
    border-radius: 12px;
    box-shadow: 0 0 10px rgba(0,0,0,0.1);
}
.token-box {
    display: inline-block;
    padding: 6px 10px;
    margin: 4px;
    border-radius: 6px;
    background-color: #e9ecef;
    font-size: 15px;
    color: #212529;
    border: 1px solid #dee2e6;
}
</style>
""", unsafe_allow_html=True)

st.image("./urdugpt.png", width=120)
st.title("🌐 UrduGPT Translator")
st.markdown("Your friendly LLM-powered translator from English → Urdu ✨")

with st.sidebar:
    st.header("⚙️ Settings")
    view_mode = st.radio("Output View", ["Sentence", "Token-by-Token"])

if "model" not in st.session_state:
    with st.spinner("Loading UrduGPT model..."):
        st.session_state.model = load_model()

if "history" not in st.session_state:
    st.session_state.history = []

text_input = st.text_area("✏️ Type English text:", height=100)
if st.button("Translate") and text_input.strip():
    with st.spinner("Translating..."):
        ids, toks, confs = decode_with_confidence(st.session_state.model, text_input.strip())
        urdu_out = tokenizer_ur.decode(ids)
        st.session_state.history.append((text_input.strip(), urdu_out))

    st.success("✅ Translation complete")
    if view_mode == "Sentence":
        st.markdown(f"**🇵🇰 UrduGPT says:**\n\n{urdu_out}")
    else:
        st.markdown("**🧠 Token-level confidence:**")
        for tok, conf in zip(toks, confs):
            st.markdown(f'<span class="token-box">{tok} ({conf:.2f})</span>', unsafe_allow_html=True)

if st.session_state.history:
    st.divider()
    st.subheader("🕘 Translation History")
    hist_df = pd.DataFrame(st.session_state.history, columns=["English", "UrduGPT"])
    st.dataframe(hist_df)

    csv_data = convert_df_to_csv(hist_df)
    st.download_button("⬇️ Download CSV", csv_data, "urdugpt_history.csv", "text/csv")

    pdf_data = export_to_pdf(st.session_state.history)
    st.download_button("⬇️ Download PDF", pdf_data, "urdugpt_history.pdf")

    docx_data = export_to_docx(st.session_state.history)
    st.download_button("⬇️ Download Word", docx_data, "urdugpt_history.docx")
